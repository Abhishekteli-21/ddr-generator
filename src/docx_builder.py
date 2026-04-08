
import datetime
import logging
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ── Typography ─────────────────────────────────────────────────────────────
FONT_NAME = "Times New Roman"

# ── Colour palette ─────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1A, 0x3A, 0x5C)
BLUE_DARK  = RGBColor(0x1F, 0x49, 0x7D)
BLUE_MED   = RGBColor(0x2E, 0x75, 0xB6)
GREY_TEXT  = RGBColor(0x40, 0x40, 0x40)
GREY_LIGHT = RGBColor(0x70, 0x70, 0x70)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

SEV_COLORS = {
    "critical": "C00000",
    "high":     "E84040",
    "medium":   "FF9900",
    "low":      "70AD47",
}

PRI_COLORS = {
    "immediate":  "C00000",
    "short-term": "FF9900",
    "long-term":  "70AD47",
}

# Image column widths inside the document (A4 minus margins = ~16.92 cm usable)
IMG_COL_THERMAL = 2.8    # cm per column in the thermal pair table
IMG_COL_INSP    = 3.8    # cm per column in the inspection photo table (2-per-row)


# ── XML / style helpers ────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _cell_para(
    cell, text: str = "", bold=False, size=10,
    color=None, center=False, italic=False,
):
    cell.paragraphs[0].clear()
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold         = bold
    run.italic       = italic
    run.font.size    = Pt(size)
    run.font.name    = FONT_NAME
    if color:
        run.font.color.rgb = color
    return para


def _setup_page(doc):
    """A4 page with headers/footers margins."""
    s = doc.sections[0]
    s.page_height    = Cm(29.7)
    s.page_width     = Cm(21.0)
    s.left_margin    = Cm(1.8)
    s.right_margin   = Cm(1.8)
    s.top_margin     = Cm(3.8)
    s.bottom_margin  = Cm(3.0)
    s.header_distance = Cm(1.2)
    s.footer_distance = Cm(1.2)

def _add_page_number(run):
    """Adds a dynamic PAGE field to a run for Word footers."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "1"
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)

def _build_header_footer(doc):
    """Inject branded header & footer into the document natively."""
    section = doc.sections[0]
    
    # ── Header ───────────────────────────────────────────────────
    header = section.header
    for p in header.paragraphs:
        p.clear()
        
    # 1x2 table for Logo (Left) and Black Banner (Right)
    htbl = header.add_table(rows=1, cols=2, width=Cm(17.4))
    htbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    lcell = htbl.cell(0, 0)
    lcell.width = Cm(3.5)
    rcell = htbl.cell(0, 1)
    rcell.width = Cm(13.9)
    
    lcell.paragraphs[0].clear()
    lpara = lcell.paragraphs[0]
    lpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_path = Path("assets/urbanroof_logo.png")
    if logo_path.exists():
        lpara.add_run().add_picture(str(logo_path), width=Cm(2.5))
    else:
        lpara.add_run("UrbanRoof").bold = True
        
    _set_cell_bg(rcell, "000000")
    rcell.paragraphs[0].clear()
    rpara = rcell.paragraphs[0]
    rpara.paragraph_format.space_before = Pt(8)
    rpara.paragraph_format.space_after = Pt(8)
    rpara.paragraph_format.left_indent = Cm(0.5)
    
    r1 = rpara.add_run("IR-, Detailed Diagnosis Report of\n")
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.font.size = Pt(11)
    r1.font.name = "Arial"
    
    r2 = rpara.add_run("Property Inspection & Thermal Analysis")
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2.font.size = Pt(9)
    r2.font.name = "Arial"
    
    # Green line below header
    bpara = header.add_paragraph()
    bpPr = bpara._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "24")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "32CD32")
    pBdr.append(bot)
    bpPr.append(pBdr)
    bpara.paragraph_format.space_after = Pt(4)
    
    # ── Footer ───────────────────────────────────────────────────
    footer = section.footer
    for p in footer.paragraphs:
        p.clear()
        
    # Green line above footer
    fpara = footer.add_paragraph()
    fpPr = fpara._p.get_or_add_pPr()
    fpBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "24")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "32CD32")
    fpBdr.append(top)
    fpPr.append(fpBdr)
    fpara.paragraph_format.space_after = Pt(4)
    
    ftbl = footer.add_table(rows=1, cols=3, width=Cm(17.4))
    ftbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ftbl.cell(0, 0).width = Cm(5.0)
    ftbl.cell(0, 1).width = Cm(7.4)
    ftbl.cell(0, 2).width = Cm(5.0)
    
    fl = ftbl.cell(0, 0).paragraphs[0]
    fl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rl = fl.add_run("www.urbanroof.in")
    rl.font.size = Pt(10)
    rl.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    rl.underline = True
    
    fc = ftbl.cell(0, 1).paragraphs[0]
    fc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = fc.add_run("UrbanRoof Private Limited")
    rc.font.size = Pt(11)
    rc.font.color.rgb = RGBColor(0xFF, 0x99, 0x00)
    rc.italic = True
    rc.font.bold = True
    
    fr = ftbl.cell(0, 2).paragraphs[0]
    fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = fr.add_run("Page ")
    rr.font.size = Pt(11)
    rr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    rr.italic = True
    prn = fr.add_run()
    prn.font.size = Pt(12)
    prn.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    prn.font.bold = True
    _add_page_number(prn)


def _add_heading(doc, text: str, level: int = 1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12 if level == 2 else 6)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.bold       = True
    run.font.name  = FONT_NAME
    if level == 1:
        run.font.size      = Pt(14)
        run.font.color.rgb = NAVY
        run.underline      = True
    else:
        run.font.size      = Pt(12)
        run.font.color.rgb = BLUE_DARK
    return para


def _add_section_rule(doc):
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E75B6")
    pBdr.append(bot)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(4)


def _add_body(doc, text: str):
    para = doc.add_paragraph(text or "Not Available")
    para.paragraph_format.space_after = Pt(4)
    for run in para.runs:
        run.font.name      = FONT_NAME
        run.font.size      = Pt(11)
        run.font.color.rgb = GREY_TEXT
    return para


def _add_bullet(doc, text: str, bold_prefix: str = ""):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = para.add_run(f"{bold_prefix}: ")
        r.bold       = True
        r.font.name  = FONT_NAME
        r.font.size  = Pt(11)
    r2 = para.add_run(text)
    r2.font.name      = FONT_NAME
    r2.font.size      = Pt(11)
    r2.font.color.rgb = GREY_TEXT
    return para


# ── Image helpers ──────────────────────────────────────────────────────────

def _img_valid(path) -> bool:
    """Return True only when path points to a non-empty, readable file."""
    if not path:
        return False
    p = Path(path)
    return p.exists() and p.stat().st_size > 0


def _build_flat_image_list(extracted_images: dict) -> list:
    """
    Flatten per-page image dict into a globally ordered list.
    Sorted by page number then extraction index.
    Used to resolve inspection photo_refs (1-indexed).
    """
    flat = []
    for page in sorted(extracted_images.keys()):
        for img_path in extracted_images[page]:
            flat.append(Path(img_path))
    return flat


def _normalize_photo_refs(refs) -> list:
    """
    Normalise photo refs from Gemini to plain integers.
    Handles: int, str "1", str "Photo 1", str "photo_1", float 1.0, etc.
    Returns sorted deduplicated list of ints (1-indexed).
    """
    result = []
    if not refs:
        return result
    for ref in refs:
        s = str(ref).strip().lower()
        # strip any non-digit prefix like "photo", "#", "image"
        import re
        digits = re.findall(r"\d+", s)
        if digits:
            result.append(int(digits[0]))
    return sorted(set(result))


def _find_inspection_photos(merged_entry: dict, flat_insp: list) -> list:
    """
    Map inspection photo_refs to actual image paths.
    Returns list of valid Paths (empty list if none found).
    """
    if not merged_entry:
        return []
    raw_refs = merged_entry.get("inspection_photo_refs", [])
    refs     = _normalize_photo_refs(raw_refs)
    paths    = []
    for ref in refs:
        idx = ref - 1   # 1-indexed → 0-indexed
        if 0 <= idx < len(flat_insp):
            p = Path(flat_insp[idx])
            if _img_valid(p):
                paths.append(p)
            else:
                logger.debug(f"  Photo ref {ref} → index {idx}: file missing or empty")
        else:
            logger.debug(
                f"  Photo ref {ref} out of range (flat_insp has {len(flat_insp)} entries)"
            )
    return paths


def _find_thermal_images_for_page(extracted_images: dict, page_num) -> tuple:
    """
    Return (site_photo_path, thermal_heatmap_path) for a thermal PDF page.

    Within each thermal page:
      imgs[0] = thermal heatmap   (top image, false-colour)
      imgs[1] = regular site photo (bottom image)

    Falls back up to ±2 pages if the target page has no images.
    Returns (None, None) when nothing is found.
    """
    if page_num is None:
        return None, None
    page_num = int(page_num)
    for offset in [0, 1, -1, 2, -2]:
        imgs = extracted_images.get(page_num + offset, [])
        if imgs:
            heatmap    = imgs[0] if len(imgs) > 0 else None
            site_photo = imgs[1] if len(imgs) > 1 else None
            # Validate
            heatmap    = heatmap    if _img_valid(heatmap)    else None
            site_photo = site_photo if _img_valid(site_photo) else None
            if heatmap or site_photo:
                logger.debug(
                    f"  Thermal page {page_num} (offset {offset}): "
                    f"heatmap={'✓' if heatmap else '✗'}  "
                    f"site_photo={'✓' if site_photo else '✗'}"
                )
                return site_photo, heatmap   # caller names: regular_img, thermal_img
    return None, None


def _insert_image_in_para(para, img_path, width_cm: float = 5.5):
    """
    Insert image into a paragraph (centred).
    Shows italic placeholder text on failure.
    Returns True on success.
    """
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if _img_valid(img_path):
        try:
            para.add_run().add_picture(str(img_path), width=Cm(width_cm))
            return True
        except Exception as exc:
            logger.warning(f"  Image insert failed ({img_path}): {exc}")

    # Placeholder
    r = para.add_run("[ Image Not Available ]")
    r.font.name      = FONT_NAME
    r.font.size      = Pt(9)
    r.font.color.rgb = GREY_LIGHT
    r.italic         = True
    return False


def _add_header_cell(cell, text: str):
    _set_cell_bg(cell, "1F497D")
    _cell_para(cell, text, bold=True, color=WHITE, center=True, size=9)


# ── Image section builder ──────────────────────────────────────────────────

def _add_image_section(
    doc,
    insp_photos: list,      # list of Path objects from inspection PDF
    site_photo,             # regular site photo from thermal PDF
    thermal_heatmap,        # thermal heatmap from thermal PDF
):
    """
    Render images in the style of the sample DDR:

    ── Thermal Evidence ──────────────────────────────────────────────────────
    ┌──────────────────────────┬──────────────────────────┐
    │  Site Photo (Thermal)    │  Thermal Heatmap         │
    │  [image / placeholder]   │  [image / placeholder]   │
    └──────────────────────────┴──────────────────────────┘

    ── Inspection Site Evidence (shown only when photos exist) ───────────────
    ┌──────────────────────────┬──────────────────────────┐
    │  Inspection Photo 1      │  Inspection Photo 2      │
    │  [image]                 │  [image / placeholder]   │
    └──────────────────────────┴──────────────────────────┘
    (additional rows added for photo 3, 4, etc.)
    """
    has_thermal = _img_valid(site_photo) or _img_valid(thermal_heatmap)
    has_insp    = bool(insp_photos)

    if not has_thermal and not has_insp:
        p = _add_body(doc, "Images: Not Available")
        return

    # ── Thermal pair table ────────────────────────────────────────────────
    if has_thermal:
        _add_sub_label(doc, "Thermal Evidence")

        tbl = doc.add_table(rows=2, cols=2)
        tbl.style     = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Headers
        _add_header_cell(tbl.rows[0].cells[0], "Site Photo")
        _add_header_cell(tbl.rows[0].cells[1], "Thermal Heatmap")

        # Images
        img_row = tbl.rows[1].cells
        for cell in img_row:
            cell.paragraphs[0].clear()

        _insert_image_in_para(img_row[0].paragraphs[0], site_photo,       width_cm=6.8)
        _insert_image_in_para(img_row[1].paragraphs[0], thermal_heatmap,  width_cm=6.8)

        doc.add_paragraph()

    # ── Inspection photos grid ────────────────────────────────────────────
    if has_insp:
        _add_sub_label(doc, "Inspection Site Evidence")

        # Build 2-column grid; one row per pair of photos
        COLS = 2
        rows_needed = (len(insp_photos) + COLS - 1) // COLS  # ceil division
        tbl = doc.add_table(rows=rows_needed * 2, cols=COLS)  # header+image per pair
        tbl.style     = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        photo_idx = 0
        for row_pair in range(rows_needed):
            header_row = tbl.rows[row_pair * 2]
            image_row  = tbl.rows[row_pair * 2 + 1]

            for col in range(COLS):
                photo_num = row_pair * COLS + col + 1   # 1-indexed label
                hcell     = header_row.cells[col]
                icell     = image_row.cells[col]
                icell.paragraphs[0].clear()

                if photo_idx < len(insp_photos):
                    _add_header_cell(hcell, f"Inspection Photo {photo_num}")
                    _insert_image_in_para(
                        icell.paragraphs[0],
                        insp_photos[photo_idx],
                        width_cm=6.8,
                    )
                    photo_idx += 1
                else:
                    # Empty filler cell (odd number of photos)
                    _set_cell_bg(hcell, "D6E4F0")
                    _cell_para(hcell, "", size=9, center=True)
                    # leave image cell blank

        doc.add_paragraph()


def _add_sub_label(doc, text: str):
    """Small italic label above an image table."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(text)
    run.italic         = True
    run.font.name      = FONT_NAME
    run.font.size      = Pt(9)
    run.font.color.rgb = BLUE_MED


# ── Thermal readings mini-table ────────────────────────────────────────────

def _add_thermal_readings(doc, td: dict):
    if not td:
        return
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"

    headers = ["Hotspot Temp", "Coldspot Temp", "Emissivity", "Date"]
    values  = [
        str(td.get("hotspot",    td.get("hotspot_temp",  "N/A"))),
        str(td.get("coldspot",   td.get("coldspot_temp", "N/A"))),
        str(td.get("emissivity", "N/A")),
        str(td.get("date",       td.get("thermal_date",  "N/A"))),
    ]

    for i, (h, v) in enumerate(zip(headers, values)):
        _set_cell_bg(tbl.rows[0].cells[i], "D6E4F0")
        _cell_para(tbl.rows[0].cells[i], h, bold=True, size=9, center=True)
        _cell_para(tbl.rows[1].cells[i], v, size=10, center=True)

    doc.add_paragraph()


# ── Cover page ─────────────────────────────────────────────────────────────

def _build_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = banner.add_run("DETAILED DIAGNOSTIC REPORT")
    br.bold            = True
    br.font.size       = Pt(22)
    br.font.color.rgb  = NAVY
    br.font.name       = FONT_NAME

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Property Inspection & Thermal Analysis")
    sr.font.size       = Pt(14)
    sr.italic          = True
    sr.font.color.rgb  = BLUE_MED
    sr.font.name       = FONT_NAME

    doc.add_paragraph()

    # Horizontal rule
    rule = doc.add_paragraph()
    pPr  = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A3A5C")
    pBdr.append(bot)
    pPr.append(pBdr)
    doc.add_paragraph()

    # Metadata table
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    meta_rows = [
        ("Report Type",    "Detailed Diagnostic Report (DDR)"),
        ("Document Date",  datetime.date.today().strftime("%d %B %Y")),
        ("Report Version", "1.0"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        _set_cell_bg(tbl.rows[i].cells[0], "D6E4F0")
        _cell_para(tbl.rows[i].cells[0], k, bold=True, size=11)
        _cell_para(tbl.rows[i].cells[1], v, size=11)

    doc.add_paragraph()
    doc.add_paragraph()

    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disc.add_run(
        "This report is prepared based on physical inspection observations "
        "and thermal imaging data. All findings are specific to the date of inspection."
    )
    dr.font.size      = Pt(9)
    dr.italic         = True
    dr.font.color.rgb = GREY_LIGHT
    dr.font.name      = FONT_NAME


# ── Main builder ──────────────────────────────────────────────────────────

def build_docx(
    ddr_content:       dict,
    merged_data:       list,
    extracted_images:  dict,            # thermal PDF  {page: [Path …]}
    output_path:       Path,
    inspection_images: dict = None,     # inspection PDF {page: [Path …]}
):
    """
    Assemble the full DDR Word document and save to output_path.

    Parameters
    ----------
    ddr_content       : structured report dict from Gemini
    merged_data       : merged inspection + thermal entries
    extracted_images  : per-page images from the thermal PDF
    output_path       : destination .docx path
    inspection_images : per-page images from the inspection PDF (site photos)
    """
    doc = Document()
    _setup_page(doc)
    _build_header_footer(doc)

    # ── Build lookup structures ───────────────────────────────────────────
    flat_insp: list = _build_flat_image_list(inspection_images or {})
    logger.info(f"Inspection image pool: {len(flat_insp)} images for photo-ref mapping")

    merged_lookup: dict = {}
    for item in merged_data:
        key = str(item.get("area_name", "")).lower().strip()
        merged_lookup[key] = item

    # ════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════════════
    _build_cover(doc)
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 1 — Property Issue Summary
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "1.  Property Issue Summary")
    _add_section_rule(doc)
    _add_body(doc, ddr_content.get("property_issue_summary", "Not Available"))
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 2 — Area-wise Observations
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "2.  Area-wise Observations")
    _add_section_rule(doc)

    area_obs_list = ddr_content.get("area_wise_observations", [])
    for area_obs in area_obs_list:
        area_name  = area_obs.get("area_name", "Unknown Area")
        obs_text   = area_obs.get("observation_text", "Not Available")
        thermal_td = area_obs.get("thermal_data") or {}

        _add_heading(doc, area_name, level=2)
        _add_body(doc, obs_text)

        if isinstance(thermal_td, dict) and thermal_td:
            _add_thermal_readings(doc, thermal_td)

        # Resolve images
        merged_entry = merged_lookup.get(area_name.lower().strip())
        thermal_page = merged_entry.get("thermal_page") if merged_entry else None

        site_photo, thermal_heatmap = _find_thermal_images_for_page(
            extracted_images, thermal_page
        )
        insp_photos = _find_inspection_photos(merged_entry, flat_insp)

        if insp_photos:
            logger.debug(
                f"  [{area_name}] {len(insp_photos)} inspection photo(s), "
                f"thermal_page={thermal_page}, "
                f"heatmap={'✓' if thermal_heatmap else '✗'}, "
                f"site_photo={'✓' if site_photo else '✗'}"
            )
        else:
            logger.debug(
                f"  [{area_name}] no inspection photos, "
                f"thermal_page={thermal_page}, "
                f"heatmap={'✓' if thermal_heatmap else '✗'}"
            )

        _add_image_section(doc, insp_photos, site_photo, thermal_heatmap)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 3 — Probable Root Cause
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "3.  Probable Root Cause")
    _add_section_rule(doc)
    root_causes = ddr_content.get("probable_root_cause", [])
    if root_causes:
        for item in root_causes:
            _add_bullet(
                doc,
                item.get("root_cause", "Not Available"),
                bold_prefix=item.get("area_name", "Unknown"),
            )
    else:
        _add_body(doc, "Not Available")
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 4 — Severity Assessment
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "4.  Severity Assessment")
    _add_section_rule(doc)
    severity_list = ddr_content.get("severity_assessment", [])
    if severity_list:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        for cell, txt in zip(tbl.rows[0].cells, ["Area", "Severity", "Reasoning"]):
            _set_cell_bg(cell, "1F497D")
            _cell_para(cell, txt, bold=True, color=WHITE, size=11)

        for item in severity_list:
            row   = tbl.add_row()
            sev   = str(item.get("severity", "Unknown"))
            color = SEV_COLORS.get(sev.lower(), "CCCCCC")
            dark  = sev.lower() in {"low", "medium"}
            tc    = None if dark else WHITE
            _cell_para(row.cells[0], item.get("area_name",  "Unknown"),       size=10)
            _set_cell_bg(row.cells[1], color)
            _cell_para(row.cells[1], sev, bold=True, color=tc, size=10, center=True)
            _cell_para(row.cells[2], item.get("reasoning", "Not Available"),  size=10)
    else:
        _add_body(doc, "Not Available")
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 5 — Recommended Actions
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "5.  Recommended Actions")
    _add_section_rule(doc)
    rec_actions = ddr_content.get("recommended_actions", [])
    if rec_actions:
        for item in rec_actions:
            area     = item.get("area_name", "Unknown")
            priority = item.get("priority", "")
            actions  = item.get("actions", [])

            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            r1 = para.add_run(area)
            r1.bold            = True
            r1.font.color.rgb  = BLUE_DARK
            r1.font.name       = FONT_NAME
            r1.font.size       = Pt(11)

            if priority:
                r2 = para.add_run(f"   [{priority}]")
                r2.italic         = True
                r2.font.size      = Pt(10)
                r2.font.name      = FONT_NAME
                r2.font.color.rgb = GREY_LIGHT

            for action in actions:
                _add_bullet(doc, action)
    else:
        _add_body(doc, "Not Available")
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 6 — Additional Notes
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "6.  Additional Notes")
    _add_section_rule(doc)
    _add_body(doc, ddr_content.get("additional_notes", "Not Available"))

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 7 — Missing or Unclear Information
    # ════════════════════════════════════════════════════════════════════════
    _add_heading(doc, "7.  Missing or Unclear Information")
    _add_section_rule(doc)
    missing = ddr_content.get("missing_or_unclear_information", [])
    if not missing:
        _add_body(doc, "No missing or unclear information identified.")
    else:
        for item in missing:
            if isinstance(item, dict):
                desc   = item.get("item",   str(item))
                source = item.get("source", "")
                _add_bullet(doc, desc + (f"  [Affects: {source}]" if source else ""))
            else:
                _add_bullet(doc, str(item))

    # ── Save ──────────────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info(f"DDR saved → {output_path}")
    return output_path
